"""
AURA — Asistente de Redacción Radiológica Asistida por IA
Versión 1.0.0
"""

import streamlit as st
import openai
import anthropic
import tempfile
import os
import re
import json
import time
from datetime import datetime
from pathlib import Path
from docx import Document

# ─────────────────────────────────────────────
# CONFIGURACIÓN DE PÁGINA
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="AURA · Radiología",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────────────
# CSS GLOBAL — Diseño oscuro estilo gōster
# ─────────────────────────────────────────────
st.markdown("""
<style>
/* ── Tipografía y Reset ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body, .stApp {
    background: #0d0d0f !important;
    color: #e2e2e5 !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px;
}

/* ── Ocultar elementos de Streamlit ── */
#MainMenu, footer, header, .stDeployButton,
[data-testid="stToolbar"], [data-testid="stDecoration"],
[data-testid="stStatusWidget"] { display: none !important; }

.block-container {
    padding: 0 !important;
    max-width: 100% !important;
}

/* ── Layout: dos columnas sin gap visible ── */
[data-testid="stHorizontalBlock"] {
    gap: 0 !important;
    align-items: stretch !important;
}

/* ── Sidebar oculto ── */
[data-testid="stSidebar"] { display: none !important; }

/* ── Scrollbar global ── */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #2a2a30; border-radius: 2px; }

/* ── Botones Streamlit reset ── */
.stButton > button {
    border: none !important;
    outline: none !important;
    cursor: pointer;
    transition: all 0.15s ease !important;
}

/* ── Selectbox / inputs ── */
.stSelectbox > div > div,
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: #1a1a1f !important;
    border: 1px solid #2a2a35 !important;
    border-radius: 6px !important;
    color: #e2e2e5 !important;
    font-family: 'Inter', sans-serif !important;
}

.stSelectbox > div > div:hover,
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #2fe4a3 !important;
    box-shadow: 0 0 0 1px #2fe4a330 !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    background: #141418 !important;
    border: 1px solid #232328 !important;
    border-radius: 6px !important;
    color: #9999a8 !important;
    font-size: 12px !important;
}

/* ── Divider ── */
hr { border: none; border-top: 1px solid #232328; margin: 12px 0; }

/* ── Métricas ── */
[data-testid="stMetric"] {
    background: #141418;
    border: 1px solid #232328;
    border-radius: 8px;
    padding: 12px 16px !important;
}
[data-testid="stMetricLabel"] { color: #6666777 !important; font-size: 11px !important; }
[data-testid="stMetricValue"] { color: #2fe4a3 !important; font-size: 20px !important; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# CONSTANTES Y CONFIGURACIÓN
# ─────────────────────────────────────────────
MODELOS_DISPONIBLES = {
    "GPT-4o Mini": "gpt-4o-mini",
    "GPT-4.1 Mini": "gpt-4.1-mini",
    "Claude Haiku 3.5": "claude-haiku-3-5",
    "Claude Sonnet 4": "claude-sonnet-4-5",
}

SLASH_TEMPLATES = {
    "/kl": {
        "titulo": "Kellgren-Lawrence",
        "texto": "Cambios de osteoartrosis grado [GRADO] según Kellgren-Lawrence: [DESCRIPCIÓN hallazgos: pinzamiento del espacio articular, esclerosis subcondral, osteofitos marginales]."
    },
    "/birads": {
        "titulo": "BI-RADS",
        "texto": "Categoría BI-RADS [0-6]: [DESCRIPCIÓN]. [RECOMENDACIÓN]."
    },
    "/pirads": {
        "titulo": "PI-RADS",
        "texto": "Lesión en [ZONA] con puntuación PI-RADS [1-5]. [DESCRIPCIÓN características en difusión/T2/DCE]."
    },
    "/tirads": {
        "titulo": "TI-RADS",
        "texto": "Nódulo tiroideo en [LÓBULO] con puntuación TI-RADS [1-5]: [composición], [ecogenicidad], [forma], [márgenes], [focos ecogénicos]."
    },
    "/lirads": {
        "titulo": "LI-RADS",
        "texto": "Observación hepática en [SEGMENTO] categoría LI-RADS [1-5/M/TIV]: [DESCRIPCIÓN criterios mayores y auxiliares]."
    },
    "/stoller": {
        "titulo": "Clasificación Stoller (Menisco)",
        "texto": "Señal meniscal grado [0-3] según Stoller en el menisco [medial/lateral], cuerno [anterior/posterior]: [DESCRIPCIÓN]."
    },
    "/icrs": {
        "titulo": "ICRS (Cartílago)",
        "texto": "Defecto condral grado [0-4] ICRS en [LOCALIZACIÓN]: [superficie], [profundidad], [extensión estimada en mm²]."
    },
    "/spetzler": {
        "titulo": "Spetzler-Martin (MAV)",
        "texto": "MAV con puntuación Spetzler-Martin [1-5]: tamaño [<3/3-6/>6 cm], localización [elocuente/no elocuente], drenaje venoso [superficial/profundo]."
    },
    "/toast": {
        "titulo": "TOAST (ACV)",
        "texto": "Hallazgos compatibles con infarto [TERRITORIO] de probable etiología [aterotrombótica/cardioembólica/lacunar/indeterminada] según clasificación TOAST."
    },
    "/aast": {
        "titulo": "AAST (Trauma Órgano Sólido)",
        "texto": "Lesión de [ÓRGANO] grado [I-V] según escala AAST: [DESCRIPCIÓN laceración, hematoma, lesión vascular]."
    },
}

PROMPT_SISTEMA = """Eres AURA, un asistente especializado en redacción de informes radiológicos en español para radiólogos mexicanos.

REGLAS ABSOLUTAS:
1. Usa terminología radiológica española correcta y precisa
2. SIEMPRE usa "osteoartrosis" (NUNCA "osteoartritis" ni "cambios degenerativos")
3. SIEMPRE usa "desgarro" para roturas de tendones/ligamentos/meniscos (NUNCA "ruptura" ni "rotura")
4. NUNCA uses lenguaje vago como "cambios degenerativos", "leve", "moderado" sin especificar
5. Redacta en PROSA NARRATIVA continua, NUNCA en listas con bullets o guiones
6. Usa voz activa y tiempo presente
7. Incluye dimensiones, localización precisa y características específicas cuando se mencionan
8. Si el dictado menciona una clasificación, úsala correctamente con su terminología oficial
9. El informe debe tener: TÉCNICA (si aplica), HALLAZGOS y CONCLUSIÓN
10. En la CONCLUSIÓN, sintetiza de forma jerárquica: hallazgo principal primero, secundarios después
11. Evita redundancias entre HALLAZGOS y CONCLUSIÓN
12. Usa el nombre correcto de las estructuras anatómicas en nomenclatura anatómica española

FORMATO DE SALIDA:
- Texto plano estructurado con las secciones en MAYÚSCULAS seguidas de dos puntos
- Párrafos separados por salto de línea doble
- Sin markdown, sin bullets, sin numeración
"""

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
defaults = {
    "transcripcion": "",
    "informe": "",
    "historial": [],
    "plantilla_activa": None,
    "plantilla_nombre": "Sin plantilla activa",
    "modo_dictado": False,
    "audio_bytes": None,
    "modelo_seleccionado": "GPT-4o Mini",
    "contador_informes": 0,
    "tiempo_total_generacion": 0.0,
    "api_key_openai": "",
    "api_key_anthropic": "",
    "api_key_configurada": False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
def detectar_slash_command(texto: str) -> tuple[str | None, str]:
    """Detecta si el texto inicia con un slash command y retorna (comando, texto_restante)."""
    texto = texto.strip()
    for cmd in SLASH_TEMPLATES:
        if texto.lower().startswith(cmd):
            resto = texto[len(cmd):].strip()
            return cmd, resto
    return None, texto


def aplicar_slash_template(cmd: str, contexto: str = "") -> str:
    """Combina la plantilla del slash command con contexto adicional del dictado."""
    tmpl = SLASH_TEMPLATES[cmd]
    if contexto:
        return f"Usando la clasificación {tmpl['titulo']}, redacta un hallazgo basado en:\n{contexto}\n\nPlantilla base: {tmpl['texto']}"
    return f"Redacta un informe usando la clasificación {tmpl['titulo']}. Plantilla: {tmpl['texto']}"


def cargar_plantilla_docx(archivo) -> str:
    """Extrae texto de un archivo .docx subido."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(archivo.read())
        tmp_path = tmp.name
    doc = Document(tmp_path)
    os.unlink(tmp_path)
    partes = []
    for para in doc.paragraphs:
        if para.text.strip():
            partes.append(para.text.strip())
    for tabla in doc.tables:
        for fila in tabla.rows:
            fila_texto = " | ".join(c.text.strip() for c in fila.cells if c.text.strip())
            if fila_texto:
                partes.append(fila_texto)
    return "\n".join(partes)


def transcribir_audio(audio_bytes: bytes, api_key: str) -> str:
    """Transcribe audio con Whisper via OpenAI."""
    client = openai.OpenAI(api_key=api_key)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
        tmp.write(audio_bytes)
        tmp_path = tmp.name
    try:
        with open(tmp_path, "rb") as f:
            result = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                language="es",
                prompt="Informe radiológico. Términos: osteoartrosis, desgarro, hiperintensidad, hipointensidad, cortical, subcortical, menisco, ligamento cruzado, manguito rotador, BI-RADS, PI-RADS.",
            )
        return result.text
    finally:
        os.unlink(tmp_path)


def generar_informe_openai(dictado: str, modelo: str, api_key: str, plantilla: str = "") -> str:
    """Genera informe con modelos OpenAI."""
    client = openai.OpenAI(api_key=api_key)
    user_content = dictado
    if plantilla:
        user_content = f"PLANTILLA DE REFERENCIA:\n{plantilla}\n\nDICTADO DEL RADIÓLOGO:\n{dictado}"

    response = client.chat.completions.create(
        model=modelo,
        messages=[
            {"role": "system", "content": PROMPT_SISTEMA},
            {"role": "user", "content": user_content},
        ],
        temperature=0.3,
        max_tokens=1800,
    )
    return response.choices[0].message.content.strip()


def generar_informe_anthropic(dictado: str, modelo: str, api_key: str, plantilla: str = "") -> str:
    """Genera informe con modelos Anthropic Claude."""
    client = anthropic.Anthropic(api_key=api_key)
    user_content = dictado
    if plantilla:
        user_content = f"PLANTILLA DE REFERENCIA:\n{plantilla}\n\nDICTADO DEL RADIÓLOGO:\n{dictado}"

    message = client.messages.create(
        model=modelo,
        max_tokens=1800,
        system=PROMPT_SISTEMA,
        messages=[{"role": "user", "content": user_content}],
    )
    return message.content[0].text.strip()


def generar_informe(dictado: str, modelo_nombre: str, plantilla: str = "") -> str:
    """Dispatcher para elegir el proveedor correcto según el modelo."""
    modelo_id = MODELOS_DISPONIBLES[modelo_nombre]

    # Detectar slash command en el dictado
    cmd, texto_restante = detectar_slash_command(dictado)
    if cmd:
        dictado_final = aplicar_slash_template(cmd, texto_restante)
    else:
        dictado_final = dictado

    if modelo_id.startswith("gpt"):
        if not st.session_state.api_key_openai:
            raise ValueError("Se requiere API Key de OpenAI para este modelo.")
        return generar_informe_openai(dictado_final, modelo_id, st.session_state.api_key_openai, plantilla)
    elif modelo_id.startswith("claude"):
        if not st.session_state.api_key_anthropic:
            raise ValueError("Se requiere API Key de Anthropic para este modelo.")
        return generar_informe_anthropic(dictado_final, modelo_id, st.session_state.api_key_anthropic, plantilla)
    else:
        raise ValueError(f"Modelo no reconocido: {modelo_id}")


def guardar_en_historial(transcripcion: str, informe: str, modelo: str):
    """Añade entrada al historial de sesión."""
    entrada = {
        "id": st.session_state.contador_informes + 1,
        "timestamp": datetime.now().strftime("%H:%M"),
        "modelo": modelo,
        "transcripcion_preview": transcripcion[:80] + "..." if len(transcripcion) > 80 else transcripcion,
        "informe": informe,
    }
    st.session_state.historial.insert(0, entrada)
    st.session_state.contador_informes += 1


def exportar_docx(informe: str) -> bytes:
    """Genera un archivo .docx con el informe formateado."""
    doc = Document()

    # Configurar márgenes
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # Encabezado
    titulo = doc.add_heading("INFORME RADIOLÓGICO", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("_" * 60)
    doc.add_paragraph()

    # Contenido del informe
    for linea in informe.split("\n"):
        if linea.strip():
            p = doc.add_paragraph(linea)
            p.paragraph_format.space_after = Pt(6)

    # Pie
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("Generado con AURA · Asistente de Redacción Radiológica")

    # Guardar en bytes
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    with open(tmp_path, "rb") as f:
        data = f.read()
    os.unlink(tmp_path)
    return data


# ─────────────────────────────────────────────
# BARRA SUPERIOR
# ─────────────────────────────────────────────
st.markdown("""
<div style="
    background: #101014;
    border-bottom: 1px solid #1e1e24;
    padding: 10px 20px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    position: sticky;
    top: 0;
    z-index: 999;
">
    <div style="display:flex; align-items:center; gap:10px;">
        <span style="font-size:18px; font-weight:600; letter-spacing:-0.3px; color:#fff;">
            ⚡ AURA
        </span>
        <span style="font-size:11px; color:#555; font-weight:400; letter-spacing:0.5px; text-transform:uppercase;">
            Radiología · IA
        </span>
    </div>
    <div style="display:flex; gap:8px; align-items:center;">
""", unsafe_allow_html=True)

# Badges de estado en la barra superior
col_b1, col_b2, col_b3 = st.columns([1, 1, 1])
with col_b1:
    plantilla_badge = st.session_state.plantilla_nombre
    color_b1 = "#1a3a2a" if "Sin" not in plantilla_badge else "#1a1a20"
    st.markdown(f"""
    <div style="background:{color_b1}; border:1px solid #2a2a35; border-radius:20px;
         padding:3px 12px; font-size:11px; color:#9999a8; white-space:nowrap; text-align:center;">
        📄 {plantilla_badge}
    </div>
    """, unsafe_allow_html=True)
with col_b2:
    modelo_badge = st.session_state.modelo_seleccionado
    st.markdown(f"""
    <div style="background:#1a1a20; border:1px solid #2a2a35; border-radius:20px;
         padding:3px 12px; font-size:11px; color:#9999a8; white-space:nowrap; text-align:center;">
        🤖 {modelo_badge}
    </div>
    """, unsafe_allow_html=True)
with col_b3:
    n_inf = st.session_state.contador_informes
    st.markdown(f"""
    <div style="background:#1a1a20; border:1px solid #2a2a35; border-radius:20px;
         padding:3px 12px; font-size:11px; color:#9999a8; white-space:nowrap; text-align:center;">
        ✅ {n_inf} informe{'s' if n_inf != 1 else ''}
    </div>
    """, unsafe_allow_html=True)

st.markdown("</div></div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# LAYOUT PRINCIPAL: DOS COLUMNAS
# ─────────────────────────────────────────────
col_izq, col_der = st.columns([1, 3], gap="small")

# ════════════════════════════════════════════
# COLUMNA IZQUIERDA — Panel de Dictado
# ════════════════════════════════════════════
with col_izq:
    st.markdown("""
    <div style="
        background: #0f0f13;
        border-right: 1px solid #1a1a20;
        min-height: calc(100vh - 52px);
        padding: 20px 16px;
        display: flex;
        flex-direction: column;
        gap: 16px;
    ">
    """, unsafe_allow_html=True)

    # ── Título del panel ──
    st.markdown("""
    <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:4px;">
        <span style="font-size:11px; font-weight:600; color:#555; letter-spacing:1px; text-transform:uppercase;">
            TRANSCRIPCIÓN
        </span>
    </div>
    """, unsafe_allow_html=True)

    # ── Selector Modo: Dictado / Teclado ──
    modo = st.radio(
        "Modo de entrada",
        ["🎙 Dictado", "⌨ Teclado"],
        horizontal=True,
        label_visibility="collapsed",
        key="modo_entrada",
    )

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── MODO DICTADO ──
    if modo == "🎙 Dictado":
        st.markdown("""
        <div style="text-align:center; padding:20px 0 10px;">
            <div style="
                width:90px; height:90px; border-radius:50%;
                background: radial-gradient(circle at 35% 35%, #2fe4a3, #0d9e6e);
                box-shadow: 0 0 30px #2fe4a340, 0 0 60px #2fe4a315;
                display:inline-flex; align-items:center; justify-content:center;
                font-size:36px; cursor:pointer;
                animation: pulse 2.5s ease-in-out infinite;
            ">🎙</div>
        </div>
        <style>
        @keyframes pulse {
            0%, 100% { box-shadow: 0 0 30px #2fe4a340, 0 0 60px #2fe4a315; }
            50% { box-shadow: 0 0 45px #2fe4a360, 0 0 80px #2fe4a325; }
        }
        </style>
        <p style="text-align:center; font-size:13px; color:#777; margin-top:8px;">
            Sube un archivo de audio
        </p>
        <p style="text-align:center; font-size:11px; color:#444; margin-top:4px;">
            <b style="color:#2fe4a3;">AURA</b> interpretará el audio al generar el informe
        </p>
        """, unsafe_allow_html=True)

        audio_file = st.file_uploader(
            "Archivo de audio",
            type=["mp3", "wav", "m4a", "webm", "ogg", "flac"],
            label_visibility="collapsed",
            key="uploader_audio",
        )
        if audio_file:
            st.session_state.audio_bytes = audio_file.read()
            st.markdown(f"""
            <div style="background:#1a2e20; border:1px solid #2a4a30; border-radius:6px;
                 padding:8px 12px; font-size:12px; color:#2fe4a3; margin-top:8px;">
                ✓ {audio_file.name}
            </div>
            """, unsafe_allow_html=True)

    # ── MODO TECLADO ──
    else:
        st.markdown("""
        <p style="font-size:11px; color:#555; margin-bottom:6px;">
            Escribe o pega los hallazgos. Usa <code style="color:#2fe4a3;">/comando</code> al inicio para clasificaciones.
        </p>
        """, unsafe_allow_html=True)

        transcripcion_input = st.text_area(
            "Hallazgos",
            value=st.session_state.transcripcion,
            height=260,
            placeholder="Ej: /kl rodilla derecha con pinzamiento medial...\n\nO describe libremente los hallazgos...",
            label_visibility="collapsed",
            key="texto_manual",
        )
        st.session_state.transcripcion = transcripcion_input

        # Referencia de slash commands
        with st.expander("📎 Comandos de clasificación"):
            for cmd, data in SLASH_TEMPLATES.items():
                st.markdown(f"""
                <div style="display:flex; gap:8px; align-items:center; padding:4px 0;
                     border-bottom:1px solid #1a1a20;">
                    <code style="color:#2fe4a3; font-size:11px; min-width:70px;">{cmd}</code>
                    <span style="color:#777; font-size:11px;">{data['titulo']}</span>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── Configuración rápida ──
    st.markdown("""
    <p style="font-size:11px; color:#555; letter-spacing:0.5px; text-transform:uppercase; margin-bottom:6px;">
        Modelo IA
    </p>
    """, unsafe_allow_html=True)

    modelo_sel = st.selectbox(
        "Modelo",
        options=list(MODELOS_DISPONIBLES.keys()),
        index=list(MODELOS_DISPONIBLES.keys()).index(st.session_state.modelo_seleccionado),
        label_visibility="collapsed",
        key="selector_modelo",
    )
    st.session_state.modelo_seleccionado = modelo_sel

    # ── Plantilla DOCX ──
    st.markdown("""
    <p style="font-size:11px; color:#555; letter-spacing:0.5px; text-transform:uppercase;
       margin-top:12px; margin-bottom:6px;">
        Plantilla (.docx)
    </p>
    """, unsafe_allow_html=True)

    plantilla_file = st.file_uploader(
        "Plantilla DOCX",
        type=["docx"],
        label_visibility="collapsed",
        key="uploader_plantilla",
    )
    if plantilla_file:
        try:
            texto_plantilla = cargar_plantilla_docx(plantilla_file)
            st.session_state.plantilla_activa = texto_plantilla
            st.session_state.plantilla_nombre = plantilla_file.name[:20]
            st.success(f"✓ Plantilla cargada")
        except Exception as e:
            st.error(f"Error al leer plantilla: {e}")

    if st.session_state.plantilla_activa:
        if st.button("✕ Quitar plantilla", use_container_width=True, key="btn_quitar_plantilla"):
            st.session_state.plantilla_activa = None
            st.session_state.plantilla_nombre = "Sin plantilla activa"
            st.rerun()

    st.markdown("<hr>", unsafe_allow_html=True)

    # ── BOTÓN GENERAR INFORME ──
    st.markdown("""
    <style>
    div[data-testid="stButton"] button[kind="primary"] {
        background: linear-gradient(135deg, #2fe4a3, #0d9e6e) !important;
        color: #051a10 !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        border-radius: 8px !important;
        height: 48px !important;
        width: 100% !important;
        letter-spacing: 0.2px !important;
    }
    div[data-testid="stButton"] button[kind="primary"]:hover {
        filter: brightness(1.1) !important;
        transform: translateY(-1px) !important;
    }
    </style>
    """, unsafe_allow_html=True)

    generar_pressed = st.button(
        "⚡ Generar Informe",
        type="primary",
        use_container_width=True,
        key="btn_generar",
    )

    st.markdown("</div>", unsafe_allow_html=True)


# ════════════════════════════════════════════
# COLUMNA DERECHA — Editor de Informe
# ════════════════════════════════════════════
with col_der:
    st.markdown("""
    <div style="
        background: #0d0d0f;
        min-height: calc(100vh - 52px);
        padding: 0;
        display: flex;
        flex-direction: column;
    ">
    """, unsafe_allow_html=True)

    # ── Tabs: Editor / Historial / Configuración ──
    tab_editor, tab_historial, tab_config = st.tabs(["📝 Informe", "📋 Historial", "⚙ Configuración"])

    # ─────────────────────────────
    # TAB: EDITOR
    # ─────────────────────────────
    with tab_editor:

        # Barra de herramientas del editor
        st.markdown("""
        <div style="
            background: #111115;
            border-bottom: 1px solid #1e1e24;
            padding: 8px 16px;
            display: flex;
            align-items: center;
            gap: 4px;
            flex-wrap: wrap;
        ">
            <span style="font-size:11px; color:#444; letter-spacing:0.5px; text-transform:uppercase; margin-right:8px;">
                INFORME
            </span>
        </div>
        """, unsafe_allow_html=True)

        # ── Lógica de generación ──
        if generar_pressed:
            # Validar API keys antes
            modelo_id = MODELOS_DISPONIBLES[st.session_state.modelo_seleccionado]
            api_ok = True

            if modelo_id.startswith("gpt") and not st.session_state.api_key_openai:
                st.error("⚠ Configura tu API Key de OpenAI en la pestaña ⚙ Configuración.")
                api_ok = False
            elif modelo_id.startswith("claude") and not st.session_state.api_key_anthropic:
                st.error("⚠ Configura tu API Key de Anthropic en la pestaña ⚙ Configuración.")
                api_ok = False

            if api_ok:
                # Obtener transcripción (audio o teclado)
                texto_final = st.session_state.transcripcion

                if modo == "🎙 Dictado" and st.session_state.audio_bytes:
                    with st.spinner("🎙 Transcribiendo audio con Whisper..."):
                        try:
                            texto_final = transcribir_audio(
                                st.session_state.audio_bytes,
                                st.session_state.api_key_openai,
                            )
                            st.session_state.transcripcion = texto_final
                        except Exception as e:
                            st.error(f"Error en transcripción: {e}")
                            texto_final = ""

                if texto_final.strip():
                    with st.spinner("⚡ Generando informe..."):
                        try:
                            t0 = time.time()
                            informe_generado = generar_informe(
                                texto_final,
                                st.session_state.modelo_seleccionado,
                                st.session_state.plantilla_activa or "",
                            )
                            elapsed = time.time() - t0
                            st.session_state.informe = informe_generado
                            st.session_state.tiempo_total_generacion += elapsed
                            guardar_en_historial(
                                texto_final,
                                informe_generado,
                                st.session_state.modelo_seleccionado,
                            )
                            st.toast(f"✅ Informe generado en {elapsed:.1f}s", icon="⚡")
                        except Exception as e:
                            st.error(f"Error al generar informe: {e}")
                else:
                    st.warning("Escribe los hallazgos o sube un archivo de audio primero.")

        # ── Área de edición del informe ──
        if st.session_state.informe:
            informe_editado = st.text_area(
                "Informe radiológico",
                value=st.session_state.informe,
                height=500,
                label_visibility="collapsed",
                key="editor_informe",
            )
            st.session_state.informe = informe_editado

            # ── Acciones: copiar / exportar ──
            col_a1, col_a2, col_a3 = st.columns([1, 1, 2])
            with col_a1:
                # Botón para copiar al portapapeles via JS
                st.markdown(f"""
                <button onclick="navigator.clipboard.writeText(`{informe_editado.replace('`', "'")}`)"
                    style="
                        background:#1a1a24; border:1px solid #2a2a35; border-radius:6px;
                        color:#9999a8; font-size:12px; padding:8px 16px; cursor:pointer;
                        width:100%; font-family:Inter,sans-serif;
                    ">
                    📋 Copiar
                </button>
                """, unsafe_allow_html=True)

            with col_a2:
                docx_bytes = exportar_docx(informe_editado)
                st.download_button(
                    "⬇ Descargar .docx",
                    data=docx_bytes,
                    file_name=f"informe_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="btn_docx",
                )

            with col_a3:
                if st.button("🗑 Limpiar", use_container_width=True, key="btn_limpiar"):
                    st.session_state.informe = ""
                    st.session_state.transcripcion = ""
                    st.session_state.audio_bytes = None
                    st.rerun()

            # ── Transcripción fuente (colapsable) ──
            if st.session_state.transcripcion:
                with st.expander("🎙 Transcripción fuente"):
                    st.markdown(f"""
                    <div style="
                        background:#111115; border:1px solid #1e1e24; border-radius:6px;
                        padding:12px 16px; font-size:13px; color:#777; line-height:1.6;
                        font-family:'JetBrains Mono',monospace;
                    ">
                        {st.session_state.transcripcion}
                    </div>
                    """, unsafe_allow_html=True)

        else:
            # Estado vacío
            st.markdown("""
            <div style="
                display:flex; flex-direction:column; align-items:center; justify-content:center;
                min-height:400px; text-align:center; color:#333;
            ">
                <div style="font-size:52px; margin-bottom:16px; opacity:0.3;">⚡</div>
                <div style="font-size:15px; color:#444; margin-bottom:8px;">El informe aparecerá aquí</div>
                <div style="font-size:12px; color:#333;">
                    Dicta los hallazgos y presiona Generar informe
                </div>
            </div>
            """, unsafe_allow_html=True)

    # ─────────────────────────────
    # TAB: HISTORIAL
    # ─────────────────────────────
    with tab_historial:
        if not st.session_state.historial:
            st.markdown("""
            <div style="padding:40px; text-align:center; color:#444;">
                <div style="font-size:32px; margin-bottom:12px;">📋</div>
                <div>Los informes generados en esta sesión aparecerán aquí</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            # Métricas rápidas
            col_m1, col_m2 = st.columns(2)
            with col_m1:
                st.metric("Informes en sesión", st.session_state.contador_informes)
            with col_m2:
                avg = st.session_state.tiempo_total_generacion / max(st.session_state.contador_informes, 1)
                st.metric("Tiempo promedio", f"{avg:.1f}s")

            st.markdown("<hr>", unsafe_allow_html=True)

            for entrada in st.session_state.historial:
                with st.expander(f"#{entrada['id']} · {entrada['timestamp']} · {entrada['modelo']}"):
                    st.markdown(f"""
                    <div style="font-size:12px; color:#666; margin-bottom:8px; font-style:italic;">
                        {entrada['transcripcion_preview']}
                    </div>
                    """, unsafe_allow_html=True)
                    st.text_area(
                        "Informe",
                        value=entrada['informe'],
                        height=200,
                        label_visibility="collapsed",
                        key=f"hist_{entrada['id']}",
                    )
                    col_h1, col_h2 = st.columns(2)
                    with col_h1:
                        docx_hist = exportar_docx(entrada['informe'])
                        st.download_button(
                            "⬇ .docx",
                            data=docx_hist,
                            file_name=f"informe_{entrada['id']}_{entrada['timestamp'].replace(':','')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_hist_{entrada['id']}",
                            use_container_width=True,
                        )
                    with col_h2:
                        if st.button("Cargar en editor", key=f"load_hist_{entrada['id']}", use_container_width=True):
                            st.session_state.informe = entrada['informe']
                            st.rerun()

    # ─────────────────────────────
    # TAB: CONFIGURACIÓN
    # ─────────────────────────────
    with tab_config:
        st.markdown("""
        <div style="padding:8px 0 16px;">
            <p style="font-size:13px; color:#666; line-height:1.6;">
                Configura tus API Keys para usar AURA. Las claves se almacenan únicamente en la sesión
                y no se guardan en ningún servidor.
            </p>
        </div>
        """, unsafe_allow_html=True)

        col_cfg1, col_cfg2 = st.columns(2)

        with col_cfg1:
            st.markdown("**OpenAI** — GPT-4o Mini, GPT-4.1 Mini, Whisper")
            key_oai = st.text_input(
                "OpenAI API Key",
                value=st.session_state.api_key_openai,
                type="password",
                placeholder="sk-...",
                label_visibility="collapsed",
                key="input_key_openai",
            )
            if key_oai != st.session_state.api_key_openai:
                st.session_state.api_key_openai = key_oai

            if st.session_state.api_key_openai:
                st.markdown('<div style="color:#2fe4a3; font-size:12px;">✓ Configurada</div>', unsafe_allow_html=True)

        with col_cfg2:
            st.markdown("**Anthropic** — Claude Haiku 3.5, Claude Sonnet 4")
            key_ant = st.text_input(
                "Anthropic API Key",
                value=st.session_state.api_key_anthropic,
                type="password",
                placeholder="sk-ant-...",
                label_visibility="collapsed",
                key="input_key_anthropic",
            )
            if key_ant != st.session_state.api_key_anthropic:
                st.session_state.api_key_anthropic = key_ant

            if st.session_state.api_key_anthropic:
                st.markdown('<div style="color:#2fe4a3; font-size:12px;">✓ Configurada</div>', unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("**Acerca de AURA**")
        st.markdown("""
        <div style="font-size:12px; color:#555; line-height:1.8;">
            <div>Versión: <span style="color:#888;">1.0.0</span></div>
            <div>Modelos de transcripción: <span style="color:#888;">OpenAI Whisper-1</span></div>
            <div>Modelos de generación: <span style="color:#888;">GPT-4o Mini · GPT-4.1 Mini · Claude Haiku 3.5 · Claude Sonnet 4</span></div>
            <div style="margin-top:12px; color:#444;">
                AURA es una herramienta de apoyo. El radiólogo es responsable del contenido final del informe.
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)
