import io
from docx import Document
from docx.shared import Pt

def export_to_docx(report_text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for paragraph in report_text.splitlines():
        if paragraph.strip() == "":
            doc.add_paragraph()
            continue
        if paragraph.strip().startswith("•"):
            p = doc.add_paragraph(paragraph.strip(), style="List Bullet")
        else:
            doc.add_paragraph(paragraph)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
